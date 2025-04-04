from docx.enum.section import WD_ORIENTATION, WD_SECTION_START

print ("Aaron\n\nExodus 2,16 'Aaron wird für dich zum Volk sprechen. Es ist so, als ob du durch ihn sprichst. Und er wird deine Botschaften weitergeben, so wie ein Prophet meine.'\n\n")

# Die Idee: Anhand der Daten aus Infoinput.xlsx wird ein Dokument zusammengestellt.
# Als Grundlage für das Dokument stehen mehrere docx Dokumente zur Verfügung (Bricks)
#
#
# WICHTIG  Notwendig in cmd (admin) damit es läuft
#python -m pip install --upgrade pip
#pip install os
#pip install openpyxl
#pip install docx
#pip install python-docx

import os
import datetime
import random

import openpyxl                 # to work with excel
from openpyxl import load_workbook
from docx import Document       # to work with Document
from pathlib import Path


path = "."
document = Document("."+"\\Bricks\\00_Formular\\Bestattung_Vorlage.docx")    # Ich nutze hier ne Vorlage nutzen, damit ich Seitenzahlen ergänzen kann und Querformat einbauen kann usw.
formdoc = Document()

# Vars to check ...
lname = "NACHNAME"
sname = "VORNAME"
gname = "GEBURTSNAME"
gort = "GEBURTSORT"
bekenntnis = "BEKENNTNIS"
lanschrift = "LETZTEANSCHRIFT"
sort = "STERBEORT"
gdate = "GEBURTSDATUM"
sdate = "STERBEDATUM"
sort = "STERBEORT"
lalter = "LEBENSALTER"
nangehoerig = "NAMEANGEHÖRIGE"
aangehoerig = "ADRESSEANGEHÖRIGE"
bdate = "BESTATTUNGSDATUM"
bort = "BESTATTUNGSORT"
bvers = "BIBELVERS"
apfarr = "AMTIERENDERPFARRER"
tbemerkungen = "TEXTBEMERKUNGEN"
bstatter = "BESTATTUNGSUNTERNEHMEN"
tmotiv = "TRAUERMOTIV"
lzahl = "LIEDERANZAHL"
bform = "BESTATTUNGSFORM"
penomenn = "PERSONALPRONOMENN"   #Nominativ
penomend = "PERSONALPRONOMEND"   #Dativ
penomena = "PERSONALPRONOMENA"   #Akkusativ
ponomen = "POSSESIVPRONOMEN"
bfartikeln = "BFARTIKELN"   #Bestattunsformartikel ...
bfartikela = "BFARTIKELA"
bfartikeld = "BFARTIKELD"
bfartikelg = "BFARTIKELG"

# Erstelle eine placeholder Liste, von allen placeholdern, die in den Bricks vorhanden sind um später lplaceh durch lvars zu ersetzen

lplaceh = [lname,sname,gname,gort,bekenntnis,lanschrift,sort,gdate,sdate,sort,lalter,nangehoerig,aangehoerig,bdate,bort,bvers,apfarr,tbemerkungen,bstatter,tmotiv,lzahl,bform,penomenn,penomend,penomena,ponomen,bfartikeln,bfartikela,bfartikeld,bfartikelg]


def getTheVars():       #export variables from Excel?           könnte ich wohl auch als funktion machen, so nach dem motto ... if string value in spalte ?? == var dann geh eine spalte weiter und hol dir den Wert ....
    print("Alle Infos aus Infoinput:")
    global path
    wb = load_workbook(path+"\\Infoinput.xlsx")          # am besten baue ich hier direkt ein, dass er auf die excel tabelle im gleichen Ordner schaut ....
    #name= #!/usr/bin/python
    ws = wb.active
    global lname
    lname = (ws["B3"].value)
    global sname
    sname = (ws["C3"].value)
    print("Name:            ",sname," ",lname)
    global gname
    gname = (ws["B4"].value)
    print("Geburtsname:         ",gname)
    global gort
    gort = (ws["B8"].value)
    print("Geburtsort:      ",gort)
    global bekenntnis
    bekenntnis = (ws["B9"].value)
    print("Konfession:      ",bekenntnis)
    global lanschrift
    lanschrift = (ws["B10"].value)
    print("Letzte Anschrift:",lanschrift)
    global sort
    sort = (ws["B14"].value)
    print("Sterbeort:       ",sort)
    global gdate
    gdate = datetime.datetime.strftime((ws["B6"].value), "%d/%b/%Y")        #convert date into string ... looks better -.-
    print("Geburtsdatum:    ",gdate)
    global sdate
    sdate = datetime.datetime.strftime((ws["B13"].value), "%d/%b/%Y")        #convert date into string ... looks better -.-
    print("Sterbedatum:     ",sdate)
    global lalter
    lalter = str(ws["B13"].value.year - ws["B6"].value.year - ((ws["B13"].value.month, ws["B13"].value.day) < (ws["B6"].value.month, ws["B6"].value.day)))
    print("Lebensalter:     ",lalter)
    global nangehoerig
    nangehoerig = (ws["B15"].value)
    print("Angehörige:      ",nangehoerig)
    global aangehoerig
    aangehoerig = (ws["B16"].value)
    print("Angehörige:      ",aangehoerig)
    global bdate
    bdate = datetime.datetime.strftime((ws["B19"].value), "%d/%b/%Y")        #convert date into string ... looks better -.-
    print("Tag der Bestattung:",bdate)
    global bort
    bort = (ws["B20"].value)
    print("Bestattungsort:  ",bort)
    global apfarr
    apfarr = (ws["B22"].value)
    print("Amtierender Pfarrer:",apfarr)
    global tbemerkungen
    tbemerkungen = (ws["B23"].value)
    print("Bemerkungen/Infos:",tbemerkungen)
    global bstatter
    bstatter = (ws["B24"].value)
    print("Bestatter:       ",bstatter)
    global tmotiv
    tmotiv = (ws["B27"].value)
    print("Trauermotiv:     ",tmotiv)
    global bvers
    bvers = choseRightBrick("\\Bibelvers\\",tmotiv)
    print("Bibelvers:\n",bvers,"\n")
    global lzahl
    lzahl = (ws["B31"].value)
    print("Liederanzahl:    ",lzahl)
    global bform
    bform = (ws["B21"].value)
    print("Bestattungsform:    ",bform)
    global bfartikeln
    global bfartikela
    global bfartikeld
    global bfartikelg
    if bform == "Urne":
        bfartikeln = "Die"      #Artikel Nominativ
        bfartikela = "Die"      #Akkusativ
        bfartikeld = "Der"      #Dativ
        bfartikelg = "Der"      #Genitiv
    elif bform == "Sarg":
        bfartikeln = "Der"      #Artikel Nominativ
        bfartikela = "Den"      #Akkusativ
        bfartikeld = "Dem"      #Dativ
        bfartikelg = "Des"      #Genitiv
    global penomenn
    global penomend
    global penomena
    global ponomen
    if str(ws["B5"].value) == "weiblich":
        penomenn = "Sie"
        penomend = "Ihr"
        penomena = "Sie"
        ponomen = "Ihr"
    elif str(ws["B5"].value) == "männlich":
        penomenn = "Er"
        penomend = "Ihm"
        penomena = "Ihn"
        ponomen = "Sein"


    global lvars        # creating a list of all the vars
    lvars = [lname,sname,gname,gort,bekenntnis,lanschrift,sort,gdate,sdate,sort,lalter,nangehoerig,aangehoerig,bdate,bort,bvers,apfarr,tbemerkungen,bstatter,tmotiv,lzahl,bform,penomenn,penomend,penomena,ponomen,bfartikeln,bfartikela,bfartikeld,bfartikelg]            #put new vars here ... also put them in the lplaceh list ....

    # Die Liednamen ... glaube die müssen nicht in die listen lvars und lplaceh ....
    global lname1, lname2, lname3, lname4
    lname1 = (ws["B32"].value)
    lname2 = (ws["B33"].value)
    lname3 = (ws["B34"].value)
    lname4 = (ws["B35"].value)

    global clist
    clist = ["B3","C3","B4","B5","B6","B8","B9","B10","B13","B14","B15","B16","B19","B20","B21","B23","B24"]          #hier kommen alle Excell felder rein, die am ende gecleart werden sollen. ...
    tabclear(ws)  #rausnehmen zum Testen
    wb.save("Infoinput.xlsx")

def tabclear(tab):               #ich sollte die Excell Tabelle clearen ...
    for i in clist:
        tab[i].value = ''


def paraMove(output_doc_name, paragraph):           # to keep the style
    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        output_run.bold = run.bold                                      # Run's bold data
        output_run.italic = run.italic                                  # Run's italic data
        output_run.underline = run.underline                            # Run's underline data
        output_run.font.color.rgb = run.font.color.rgb                  # Run's color data
        output_run.style.name = run.style.name                          # Run's font data
        output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment


def brickMove(doc):
    #print("Brick Moved")
    global document
    input_doc = Document(path+doc)
    for para in input_doc.paragraphs:
        paraMove(document, para)


def fillVars():
    print("\nFilling the Vars")
    lpos=-1
    for var in lplaceh:
        lpos += 1                   #check list positions
        for paragraph in document.paragraphs:
            if var in paragraph.text:
                paragraph.text = paragraph.text.replace(str(var),lvars[lpos])

def fillFormular():                             # Das Bestattungsformular wird geöffnet und mit den entsprechenden Variablen befüllt.
    print("\nFilling the Formular")
    formdoc = Document("."+"\\Bricks\\00_Formular\\Formular_Vorlage.docx")
    lpos=-1
    for var in lplaceh:
        lpos += 1                   #check list position
        for paragraph in formdoc.paragraphs:
            if var in paragraph.text:
    #            paragraph.text = paragraph.text.replace(str(var),lvars[lpos])
                inline = paragraph.runs                     # Das braucht man wohl, damit die Formatierung erhalten bleibt, durchblicke ich nicht ganz
                for i in range(len(inline)):
                    if var in inline[i].text:
                        text = inline[i].text.replace(str(var),lvars[lpos])
                        inline[i].text = text
    formdoc.save('Bestattungsformular '+lname+'.docx')
    print("\n\nFormular wurde erstellt ...")



def choseRightBrick(brickpath,parameter):       # wenn der parameter == 0 ist, dann einfach random ....
    bricklist =[]
    brickpath = Path("."+"\\Bricks\\"+brickpath)
    for item in brickpath.iterdir():            # abfrage ob er ne excel tabelle durchgehen soll, oder eben verschiedene docx dateien
        if item.suffix == ".xlsx":
            lpos = 4                            #startwer ab dem es in der Excell tabelle einträge gibt
            excel = load_workbook(item)         #Workbook
            sheet = excel.active                #Aktive Tabellenliste/Seite
            for cell in sheet["B5":"B17"]:       #durchsuche die Spalte .. länge muss noch manuell angepasst werden
                lpos +=1                        #trakt die position in der Liste ... in diesem Fall die Zeile
                if parameter in cell[0].value:
                    bricklist.append(sheet["C"+str(lpos)].value)
        if item.suffix == ".docx":
            if parameter == 0:
                bricklist.append("\\"+str(item))
            else:
                brickdoc = Document(item)
                header = brickdoc.sections[0].header
                if parameter in header.paragraphs[0].text:
                    bricklist.append("\\"+str(item))
                #    motiv = item.section[0]
                #    header = document.sections[0].header
                #    header.paragraphs[0].text = "Trauerfeier "+sname+" "+lname
    brick = random.choice(bricklist)
    return brick



def buildEverything():
    buildIntro()
    buildAnsprache()
    buildOutro()
    buildAmGrab()


def buildIntro():
    print("\nBuilding the intro now ...")
    global path
    #Headline
    header = document.sections[0].header
    header.paragraphs[0].text = "Trauerfeier "+sname+" "+lname + "        "+bdate

    #Header
    document.add_heading("Trauerfeier von "+sname+" "+lname, 0)

    #Votum
    document.add_heading("Votum", 1)
    brickMove(choseRightBrick("Votum",0))           # ich muss die datei weglassen denk ich und die Datei mit den Random sachen auswählen????

    #Begrüßung
    document.add_heading("Begrüßung", 1)
    brickMove(choseRightBrick("\\Begrüßung\\",0))

    #Psalm
    document.add_heading("Psalm", 1)
    brickMove(choseRightBrick("\\Psalm\\",tmotiv))

    #Mögliches Lied
    if lzahl >= 3:
        document.add_heading("Lied: "+lname1, 1)

    #Eingangsgebet
    document.add_heading("Eingangsgebet", 1)
    brickMove(choseRightBrick("\\Eingangsgebet\\",tmotiv))

    #Schriftlesung
    document.add_heading("Schriftlesung", 1)
    brickMove(choseRightBrick("\\Schriftlesung\\",tmotiv))

    #Mögliches Lied
    if lzahl >= 3:
        document.add_heading("Lied: "+lname2, 1)
    else:
        document.add_heading("Lied: "+lname1, 1)
    print("Intro finished")


def buildAnsprache():
    print("\nBuilding the Ansprache now ...")
    global path

    #Traueransprache
    document.add_heading("Traueransprache", 1)
    brickMove(choseRightBrick("\\Ansprache\\"+str(tmotiv)+"\\Hinführung\\",0))
    document.add_paragraph()
    # Persönliche Highlights in Form des Motives ... Stationen aufm Weg/Säulen, die Tragen/Puzzle, die ein Mosaik ergeben
    document.add_paragraph()
    document.add_paragraph()
    brickMove(choseRightBrick("\\Ansprache\\"+str(tmotiv)+"\\Ausblick\\",0))
    print("Ansprache finished")

def buildOutro():
    print("\nBuilding the outro now ...")
    global path

    #Mögliches Lied
    if lzahl >= 3:
        document.add_heading("Lied: "+lname3, 1)
    else:
        document.add_heading("Lied: "+lname2, 1)

    #Fürbitten
    document.add_heading("Fürbitten", 1)
    brickMove(choseRightBrick("\\Fürbitten\\",0))

    #Abschiedswort
    document.add_heading("Abschiedswort", 1)
    brickMove(choseRightBrick("\\Abschiedswort\\",0))

    #Aussegnung
    document.add_heading("Aussegnung", 1)
    brickMove(choseRightBrick("\\Aussegnung\\",0))

    #Mögliches Lied
    if lzahl == 4:
        document.add_heading("Lied: "+lname4, 1)

    #Geleitwort
    document.add_heading("Geleitwort", 1)
    brickMove(choseRightBrick("\\Geleitwort\\",0))

    print("Outro finished")

def buildAmGrab():
    print("\nBuilding AmGrab ...")
    global path

    #Bestattungswort
    document.add_heading("Bestattungswort", 1)
    brickMove(choseRightBrick("\\Bestattungswort\\",0))

    #Auferstehungswort
    document.add_heading("Auferstehungswort", 1)
    brickMove(choseRightBrick("\\Auferstehungswort\\",tmotiv))

    #Vaterunser
    document.add_heading("Vaterunser", 1)
    brickMove(choseRightBrick("\\Vaterunser\\",0))

    #Segen
    document.add_heading("Segen", 1)
    brickMove(choseRightBrick("\\Segen\\",0))

    print("Am Grab finished")



try:
    getTheVars()                        # get vars from the Infoinput
    buildEverything()                   # build the doc
    fillVars()                          # filling the doc with the right vars (placeholder into vars)
    document.save('Bestattung '+sname+' '+lname+'.docx')        # just save the final doc
    print("\n\nBeerdigung wurde erstellt ...")
    fillFormular()

finally:
    print("\nBis zum nächsten Mal ....")




# ------ ToDo:  --------------------------------------------------------------------------------------------
#   x  Platzhalter ersetzen im Abschlusstext ... hoffe das klappt
#   x  Möglichkeit zur Random auswahl von Bricks, bzw. zur geordneten Auswahl
#   x  Alle Texte mit den Motiven verknüpfen, für den maximalen roten Faden ... (es muss aber die möglichkeit geben auch allgemeine Texte zu mischen ... nicht jede Begrüßung braucht gleich nen Thema)
#   x  Kompletten Ablauf erstellen
#   o  Ablauf anpassen an Bestattungsform/Ort
#   o  Groß/Kleinschreibung bei Pronomen checken ...
#   x  Ausgabedatei umbenennen in "Beerdigung xyz"
#   x  Excel Tabelle clearen am Ende oder: GUI
#   o  TODESART in Infoinput einbauen ("Er ist friedlich eingeschlafen. / Er verstarb plötzlich und unvermittelt.") Also ein Satz, der die Art und Weise des Todes beschreibt und dann eingesetzt wird an entsprechender Stelle.
#   x  Noch stärker das Regenbogenmotiv einbauen. Auch an anderer Stelle
#   o  Motivliste: Ich habe:
#           2x Weg
#           1x Fluss/Bootsfahrt (hat den längsten Text aktuell ca. 5 Minuten allein durch Hinführung und Ausblick)
#           1x Psalm 23 (at den längsten Text aktuell ca. 5 Minuten allein durch Hinführung und Ausblick)
#           1x Bild
#           1x Regenbogen/Mosaik
#           1x Hand
#           0x Säulen   (muss ich noch dringend schreiben: Verschiedene Säulen = das Leben ... was hat das Leben ausgemacht ... was hat die Person getragen, was war stabil ...)
#
#   o  krasse Beerdigung einbauen als Thema (also Suizid oder Kind ... mit besonderen Gebeten, Ansprache muss man dann schauen ...)
#
#
#   x  Lieder anpassen und ggf Ablauf anpassen
#   x  in männliche und weibliche Anrede unterscheiden
#   x  Weitere Texte erstellen ...
#   x  Formatierung der Bricks überarbeiten, die meisten sehen hässlich aus ....
#   x  Artikel bei Bestattungsform als Placeholder einbauen ....
#   x  Formatierung überarbeiten, sodass man 2 spalten Querformat hat, um direkt auszudrucken und Seitenanzahl
#   x  Beerdigungstag in die Datei übernehmen also in den Titel (habs im Header eingefügt)
#   o  Bestattungsagende der UEK hat noch ein paar ziemlich coole Gebete ...
#           Die Agende ist übertrieben gut, da gibt Gebete und coole Texte, vor allem auch gute Alternativen zu dem "standart" kram.
#           Ich muss überlegen, ob ich die einfach einstreue, die anderen ersetze, oder gezielt die neuen ansteuerbar mache ...
